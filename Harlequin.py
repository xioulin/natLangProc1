import nltk

import docx

doc = docx.Document('romantic.docx')

def getFullText(file):
    doc1 = docx.Document(file)
    fullText = []
    for para in doc1.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


print(getFullText('randomStuff/romantic.docx'))
# print(getFullText('demo.docx'))
#
# print(len(set(doc)))
# print(doc.paragraphs[2].text)
# print(doc.paragraphs[2].runs[1].text)
# print(len(doc.paragraphs[2].runs))

# print(len(doc.paragraphs))